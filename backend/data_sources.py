# data_sources.py
"""
Multi-Source Data Ingestion System
Supports CSV, DOCX, and TXT files for different business data categories
"""

import csv
import io
import os
import re
from typing import Dict, Any, List, Optional
from datetime import datetime

# Try to import docx library for Word document parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================
# DATA SOURCE CATEGORIES
# ============================================

DATA_CATEGORIES = {
    "revenue": {
        "name": "Revenue & Financials",
        "icon": "💰",
        "description": "Revenue, costs, profit, pricing data",
        "supported_formats": ["csv", "xlsx"],
        "key_metrics": ["revenue", "profit", "costs", "price", "margin", "growth"]
    },
    "customers": {
        "name": "Customer Data",
        "icon": "👥",
        "description": "Customer counts, segments, demographics",
        "supported_formats": ["csv", "xlsx"],
        "key_metrics": ["customers", "new_customers", "churned_customers", "lifetime_value"]
    },
    "reviews": {
        "name": "Customer Reviews & Feedback",
        "icon": "💬",
        "description": "Customer reviews, support tickets, feedback",
        "supported_formats": ["csv", "docx", "txt"],
        "key_metrics": ["sentiment", "satisfaction", "nps", "rating"]
    },
    "marketing": {
        "name": "Marketing Campaigns",
        "icon": "📢",
        "description": "Campaign performance, ad spend, conversions",
        "supported_formats": ["csv", "xlsx"],
        "key_metrics": ["marketing_spend", "impressions", "clicks", "conversions", "cac"]
    },
    "operations": {
        "name": "Operations & Delivery",
        "icon": "🚚",
        "description": "Delivery times, fulfillment, logistics",
        "supported_formats": ["csv", "xlsx"],
        "key_metrics": ["delivery_delay", "fulfillment_rate", "inventory", "returns"]
    },
    "general": {
        "name": "General Data",
        "icon": "📊",
        "description": "Any other business metrics",
        "supported_formats": ["csv", "xlsx", "docx", "txt"],
        "key_metrics": []
    }
}


class DataSourceManager:
    """
    Manages multiple data sources for the Digital Twin
    """
    
    def __init__(self):
        self.sources: Dict[str, Dict[str, Any]] = {}
        self.combined_state: Dict[str, float] = {}
        self.text_data: Dict[str, List[str]] = {}  # For reviews/feedback text
        self.upload_history: List[Dict[str, Any]] = []
    
    def get_categories(self) -> Dict[str, Any]:
        """Return available data categories"""
        return DATA_CATEGORIES
    
    def parse_csv(self, content: bytes, category: str) -> Dict[str, Any]:
        """Parse CSV file and extract numeric metrics"""
        try:
            decoded = content.decode("utf-8")
            reader = csv.DictReader(io.StringIO(decoded))
            rows = list(reader)
            
            if not rows:
                return {"success": False, "error": "CSV file is empty"}
            
            # Extract metrics
            metrics = {}
            for row in rows:
                for key, value in row.items():
                    clean_key = self._normalize_key(key)
                    try:
                        # Clean and parse numeric values
                        clean_value = str(value).replace(",", "").replace("$", "").replace("%", "").strip()
                        num_value = float(clean_value)
                        
                        if clean_key not in metrics:
                            metrics[clean_key] = []
                        metrics[clean_key].append(num_value)
                    except (ValueError, AttributeError):
                        pass
            
            # Calculate averages/latest for each metric
            processed = {}
            for key, values in metrics.items():
                processed[key] = {
                    "current": values[-1] if values else 0,
                    "average": sum(values) / len(values) if values else 0,
                    "min": min(values) if values else 0,
                    "max": max(values) if values else 0,
                    "count": len(values)
                }
            
            return {
                "success": True,
                "type": "csv",
                "category": category,
                "metrics": processed,
                "rows_processed": len(rows),
                "columns": list(rows[0].keys()) if rows else []
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def parse_docx(self, content: bytes, category: str) -> Dict[str, Any]:
        """Parse DOCX file and extract text content"""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx library not installed. Run: pip install python-docx"}
        
        try:
            doc = Document(io.BytesIO(content))
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables if any
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_text)
            
            # Simple sentiment analysis for reviews
            sentiment_score = self._analyze_sentiment(paragraphs)
            
            return {
                "success": True,
                "type": "docx",
                "category": category,
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "tables": table_data,
                "sentiment_score": sentiment_score,
                "word_count": sum(len(p.split()) for p in paragraphs)
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def parse_txt(self, content: bytes, category: str) -> Dict[str, Any]:
        """Parse TXT file"""
        try:
            decoded = content.decode("utf-8")
            lines = [line.strip() for line in decoded.split("\n") if line.strip()]
            
            sentiment_score = self._analyze_sentiment(lines)
            
            return {
                "success": True,
                "type": "txt",
                "category": category,
                "lines": lines,
                "line_count": len(lines),
                "sentiment_score": sentiment_score,
                "word_count": sum(len(line.split()) for line in lines)
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _normalize_key(self, key: str) -> str:
        """Normalize metric key names"""
        return key.strip().lower().replace(" ", "_").replace("-", "_")
    
    def _analyze_sentiment(self, texts: List[str]) -> float:
        """
        Simple rule-based sentiment analysis (0-100 scale)
        For production, use a proper NLP library or API
        """
        positive_words = [
            "great", "excellent", "amazing", "good", "love", "best", "awesome",
            "fantastic", "wonderful", "happy", "satisfied", "perfect", "fast",
            "helpful", "friendly", "recommend", "quality", "reliable", "easy"
        ]
        negative_words = [
            "bad", "poor", "terrible", "awful", "hate", "worst", "disappointed",
            "slow", "broken", "issue", "problem", "error", "delay", "late",
            "rude", "unhelpful", "expensive", "waste", "never", "refund"
        ]
        
        total_words = 0
        positive_count = 0
        negative_count = 0
        
        for text in texts:
            words = text.lower().split()
            total_words += len(words)
            for word in words:
                clean_word = re.sub(r'[^\w]', '', word)
                if clean_word in positive_words:
                    positive_count += 1
                elif clean_word in negative_words:
                    negative_count += 1
        
        if positive_count + negative_count == 0:
            return 50  # Neutral
        
        sentiment = (positive_count / (positive_count + negative_count)) * 100
        return round(sentiment, 1)
    
    def add_source(
        self, 
        category: str, 
        filename: str, 
        parsed_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Add a parsed data source to the manager"""
        
        if category not in self.sources:
            self.sources[category] = {"files": [], "combined_metrics": {}}
        
        # Store file data
        self.sources[category]["files"].append({
            "filename": filename,
            "uploaded_at": datetime.now().isoformat(),
            "data": parsed_data
        })
        
        # Update combined metrics for this category
        if parsed_data.get("type") == "csv" and parsed_data.get("metrics"):
            for metric, values in parsed_data["metrics"].items():
                self.sources[category]["combined_metrics"][metric] = values["current"]
                self.combined_state[metric] = values["current"]
        
        # Handle text-based sources (reviews)
        if parsed_data.get("type") in ["docx", "txt"]:
            if category not in self.text_data:
                self.text_data[category] = []
            
            if parsed_data.get("paragraphs"):
                self.text_data[category].extend(parsed_data["paragraphs"])
            elif parsed_data.get("lines"):
                self.text_data[category].extend(parsed_data["lines"])
            
            # Add sentiment to combined state
            if parsed_data.get("sentiment_score") is not None:
                self.combined_state["sentiment"] = parsed_data["sentiment_score"]
        
        # Record in history
        self.upload_history.append({
            "category": category,
            "filename": filename,
            "timestamp": datetime.now().isoformat(),
            "type": parsed_data.get("type"),
            "success": parsed_data.get("success", False)
        })
        
        return {
            "success": True,
            "category": category,
            "filename": filename,
            "metrics_added": list(parsed_data.get("metrics", {}).keys()) if parsed_data.get("metrics") else [],
            "total_sources": len(self.sources.get(category, {}).get("files", []))
        }
    
    def get_combined_state(self) -> Dict[str, Any]:
        """Get the combined state from all sources"""
        return {
            "metrics": self.combined_state,
            "sources": {
                cat: {
                    "file_count": len(data.get("files", [])),
                    "files": [f["filename"] for f in data.get("files", [])],
                    "metrics": list(data.get("combined_metrics", {}).keys())
                }
                for cat, data in self.sources.items()
            },
            "text_data_available": list(self.text_data.keys()),
            "total_files": sum(len(d.get("files", [])) for d in self.sources.values())
        }
    
    def get_text_data(self, category: str) -> List[str]:
        """Get text data for a category (e.g., reviews)"""
        return self.text_data.get(category, [])
    
    def get_upload_history(self) -> List[Dict[str, Any]]:
        """Get upload history"""
        return self.upload_history
    
    def clear_category(self, category: str) -> bool:
        """Clear all data for a category"""
        if category in self.sources:
            # Remove metrics from combined state
            for metric in self.sources[category].get("combined_metrics", {}).keys():
                if metric in self.combined_state:
                    del self.combined_state[metric]
            
            del self.sources[category]
            
            if category in self.text_data:
                del self.text_data[category]
            
            return True
        return False
    
    def clear_all(self):
        """Clear all data sources"""
        self.sources = {}
        self.combined_state = {}
        self.text_data = {}
        self.upload_history = []


# Global instance
data_source_manager = DataSourceManager()
